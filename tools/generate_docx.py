"""基于模板docx样式，将Markdown转换为格式精确匹配的DOCX。

模板格式规格（从 docs/技术文档模板.docx 提取）：
- 字体: ASCII=Times New Roman, 东亚=宋体
- 字号: 正文12pt, 摘要14pt, 部分标题"第X部分"14pt+标题12pt, 文档标题18pt
- 行距: 1.5倍行距 (line=360, rule=auto) — 所有段落统一
- 段前/段后: 0 — 所有段落统一
- 对齐: 两端对齐 (jc=both) — 正文, 标题居中/左对齐按需
- 列表首行缩进: 0.74cm (420 twips)
- 表格: 10.5pt, 表头加粗, 有边框
"""

import os
import re
from docx import Document
from docx.shared import Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

MD_PATH = "docs/submission/技术文档_设计报告_12255_VeyraLux微睿霖光.md"
TEMPLATE_PATH = "docs/技术文档模板.docx"
OUTPUT_PATH = "docs/submission/技术文档_设计报告_12255_VeyraLux微睿霖光.docx"

ASCII_FONT = "Times New Roman"
EA_FONT = "宋体"
MONO_FONT = "Consolas"

# 模板段落间距常量（从模板XML提取）
TPL_LINE_SPACING = 1.5       # 1.5倍行距
TPL_SPACE_BEFORE = Pt(0)     # 段前0
TPL_SPACE_AFTER = Pt(0)      # 段后0
TPL_LIST_INDENT = Twips(420) # 列表首行缩进420 twips = 0.74cm

# SVG流程图文件（矢量主图 + PNG回退）
FIGURES_DIR = "docs/submission/figures"
SCREENSHOTS_DIR = "docs/submission/screenshots"
FIGURE_FILES = {
    "图1-1": os.path.join(FIGURES_DIR, "fig_1_1.png"),
    "图1-2": os.path.join(FIGURES_DIR, "fig_1_2.png"),
    "图2-1": os.path.join(FIGURES_DIR, "fig_2_1.png"),
    "图2-2": os.path.join(FIGURES_DIR, "fig_2_2.png"),
    "图3-1": os.path.join(SCREENSHOTS_DIR, "pc-web.png"),
}
FIGURE_SVG_FILES = {
    "图1-1": os.path.join(FIGURES_DIR, "fig_1_1.svg"),
    "图1-2": os.path.join(FIGURES_DIR, "fig_1_2.svg"),
    "图2-1": os.path.join(FIGURES_DIR, "fig_2_1.svg"),
    "图2-2": os.path.join(FIGURES_DIR, "fig_2_2.svg"),
}

# DrawingML命名空间
NS_A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
NS_R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
NS_ASVG = 'http://schemas.microsoft.com/office/drawing/2016/SVG/main'


def _add_svg_part(doc, svg_path):
    """将SVG文件作为Part添加到文档包中，返回relationship ID。"""
    with open(svg_path, 'rb') as f:
        svg_data = f.read()
    package = doc.part.package
    # 查找下一个可用的图片编号
    max_num = 0
    for part in package.iter_parts():
        name = str(part.partname)
        if '/word/media/image' in name:
            try:
                num = int(name.rsplit('image', 1)[1].split('.')[0])
                max_num = max(max_num, num)
            except (ValueError, IndexError):
                pass
    partname = PackURI(f'/word/media/image{max_num + 1}.svg')
    svg_part = Part(partname, 'image/svg+xml', svg_data, package)
    rId = doc.part.relate_to(svg_part, RT.IMAGE)
    return rId


def _add_svg_blip(blip_element, rId_svg):
    """在a:blip元素中添加svgBlip扩展，使Word优先使用SVG矢量渲染。"""
    extLst = etree.SubElement(blip_element, f'{{{NS_A}}}extLst')
    ext = etree.SubElement(extLst, f'{{{NS_A}}}ext')
    ext.set('uri', '{96DAC541-7B7A-43D3-8B79-37D633B846F1}')
    svgBlip = etree.SubElement(ext, f'{{{NS_ASVG}}}svgBlip')
    svgBlip.set(f'{{{NS_R}}}embed', rId_svg)


def insert_figure(doc, fig_key, caption_text=""):
    """插入SVG流程图（矢量主图 + PNG回退）及居中图注。

    Word 2016+会优先使用SVG矢量渲染，旧版Word回退到PNG。
    """
    png_path = FIGURE_FILES.get(fig_key)
    svg_path = FIGURE_SVG_FILES.get(fig_key)
    if png_path and os.path.exists(png_path):
        para = doc.add_paragraph()
        set_para_format(para, WD_ALIGN_PARAGRAPH.CENTER)
        run = para.add_run()
        # 先用PNG创建inline shape（python-docx不支持SVG直接插入）
        run.add_picture(png_path, width=Pt(425))
        # 如果SVG存在，添加svgBlip扩展使Word优先使用矢量渲染
        if svg_path and os.path.exists(svg_path):
            rId_svg = _add_svg_part(doc, svg_path)
            blip = run._element.find('.//' + qn('a:blip'))
            if blip is not None:
                _add_svg_blip(blip, rId_svg)
    # 图注
    cap = caption_text or fig_key
    para = doc.add_paragraph()
    set_para_format(para, WD_ALIGN_PARAGRAPH.CENTER)
    run = para.add_run(cap)
    set_run_font(run, 10.5, bold=False)


def set_run_font(run, size_pt, bold=False, font_ascii=ASCII_FONT, font_ea=EA_FONT):
    """设置run的字体、字号、加粗，含东亚字体。"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = font_ascii
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_ascii)
    rFonts.set(qn("w:hAnsi"), font_ascii)
    rFonts.set(qn("w:eastAsia"), font_ea)


def set_para_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_line_indent_twips=None, left_indent_pt=None):
    """统一设置段落格式：1.5倍行距、段前段后0、对齐方式。

    与模板的段落间距完全一致：
    - line=360 (1.5倍), rule=auto
    - before=0, after=0
    - jc=both (两端对齐)
    """
    pf = para.paragraph_format
    pf.line_spacing = TPL_LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = TPL_SPACE_BEFORE
    pf.space_after = TPL_SPACE_AFTER
    para.alignment = alignment
    if first_line_indent_twips is not None:
        pf.first_line_indent = first_line_indent_twips
    if left_indent_pt is not None:
        pf.left_indent = left_indent_pt


def add_formatted_para(doc, text, size_pt, bold=False,
                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       first_line_indent_twips=None):
    """添加一个完整格式的段落（字体+间距+对齐）。"""
    para = doc.add_paragraph()
    set_para_format(para, alignment, first_line_indent_twips)
    run = para.add_run(text)
    set_run_font(run, size_pt, bold)
    return para


def _set_table_borders(table):
    """手动给表格加边框。"""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def _set_cell_para_format(cell):
    """设置表格单元格段落的间距格式（与模板一致）。"""
    para = cell.paragraphs[0]
    pf = para.paragraph_format
    pf.line_spacing = TPL_LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = TPL_SPACE_BEFORE
    pf.space_after = TPL_SPACE_AFTER


def add_table(doc, rows):
    """添加格式化表格，含边框和字体。"""
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    _set_table_borders(table)
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = table.rows[i].cells[j]
                _set_cell_para_format(cell)
                para = cell.paragraphs[0]
                for r in para.runs:
                    r.text = ""
                run = para.add_run(cell_text)
                set_run_font(run, 10.5, bold=(i == 0))


def process_inline_text(para, text, size_pt, bold=False, superscript_citations=True):
    """处理行内格式（**粗体**、`代码`、正文引用标号），添加到段落。"""
    pattern = r'(\*\*[^*]+\*\*|`[^`]+`|\[\d+(?:[-,，]\d+)*\])'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            set_run_font(run, size_pt, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            set_run_font(run, size_pt - 0.5, bold=bold,
                        font_ascii=MONO_FONT, font_ea=MONO_FONT)
        elif superscript_citations and re.fullmatch(r'\[\d+(?:[-,，]\d+)*\]', part):
            run = para.add_run(part)
            set_run_font(run, size_pt - 1, bold=False)
            run.font.superscript = True
        else:
            run = para.add_run(part)
            set_run_font(run, size_pt, bold=bold)


def parse_table(lines, start_idx):
    """解析Markdown表格，返回(rows, end_idx)。"""
    rows = []
    i = start_idx
    while i < len(lines) and "|" in lines[i]:
        line = lines[i].strip()
        if line.startswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            rows.append(cells)
        i += 1
    return rows, i


def convert_md_to_docx():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document(TEMPLATE_PATH)

    # 清除模板中的所有段落和表格
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}p") or child.tag.endswith("}tbl"):
            body.remove(child)

    # 设置Normal样式默认字体
    normal_style = doc.styles["Normal"]
    normal_style.font.name = ASCII_FONT
    normal_style.font.size = Pt(12)
    rpr = normal_style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), ASCII_FONT)
    rFonts.set(qn("w:hAnsi"), ASCII_FONT)
    rFonts.set(qn("w:eastAsia"), EA_FONT)

    i = 0
    in_code_block = False
    code_lines = []
    in_references = False

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # 代码块处理（可能是ASCII框图）
        if line.strip().startswith("```"):
            if in_code_block:
                # 判断是否是图2-1的ASCII框图（含制表符/框线字符）
                ascii_art = any('┌' in cl or '│' in cl or '└' in cl or '─' in cl
                               for cl in code_lines)
                if ascii_art:
                    insert_figure(doc, "图2-1", "图2-1 系统总体框图")
                else:
                    para = doc.add_paragraph()
                    set_para_format(para, WD_ALIGN_PARAGRAPH.LEFT)
                    for cl in code_lines:
                        run = para.add_run(cl + "\n")
                        set_run_font(run, 9, bold=False,
                                    font_ascii=MONO_FONT, font_ea=MONO_FONT)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 空行 / 分隔线
        if not line.strip() or line.strip() == "---":
            i += 1
            continue

        # H1 - 文档标题 (18pt bold 居中)
        if line.startswith("# ") and not line.startswith("## "):
            title = line[2:].strip()
            add_formatted_para(doc, title, 18, bold=True,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue

        # H2 - 主要章节
        if line.startswith("## ") and not line.startswith("### "):
            title = line[3:].strip()
            if not title:
                i += 1
                continue
            in_references = title.startswith("第五部分") or title == "参考文献"

            # 作品名称 → 16pt bold
            if title == "作品名称":
                add_formatted_para(doc, title, 16, bold=True)
                i += 1
                continue

            # 摘要 → 14pt bold
            if title == "摘要":
                add_formatted_para(doc, title, 14, bold=True)
                i += 1
                continue

            # 第X部分 标题 → "第X部分" 14pt bold + 标题 12pt bold
            part_match = re.match(r'^第([一二三四五])部分\s*(.*)$', title)
            if part_match:
                part_label = f"第{part_match.group(1)}部分"
                part_title = part_match.group(2)
                para = doc.add_paragraph()
                set_para_format(para, WD_ALIGN_PARAGRAPH.LEFT)
                run1 = para.add_run(part_label + "  ")
                set_run_font(run1, 14, bold=True)
                if part_title:
                    run2 = para.add_run(part_title)
                    set_run_font(run2, 12, bold=True)
                i += 1
                continue

            # 其他H2 → 14pt bold
            add_formatted_para(doc, title, 14, bold=True)
            i += 1
            continue

        # H3 - 子节标题 (1.1, 2.1等) → 12pt bold
        if line.startswith("### ") and not line.startswith("#### "):
            title = line[4:].strip()
            if title == "特性成果":
                doc.add_page_break()
            add_formatted_para(doc, title, 12, bold=True)
            i += 1
            continue

        # H4 - 子子节标题 (2.2.1等) → 12pt bold
        if line.startswith("#### "):
            title = line[5:].strip()
            add_formatted_para(doc, title, 12, bold=True)
            i += 1
            continue

        # 表格
        if line.strip().startswith("|"):
            rows, end_idx = parse_table(lines, i)
            if rows:
                add_table(doc, rows)
            i = end_idx
            continue

        # 引用块（通常作图注，若匹配流程图占位符则插入图片）
        if line.strip().startswith(">"):
            text = line.strip()[1:].strip()
            # 匹配 **图X-X** 字样
            m = re.search(r'\*\*(图\d+-\d+)\*\*\s*(.*)', text)
            if m:
                fig_key = m.group(1)
                caption_tail = m.group(2).strip()
                if fig_key in FIGURE_FILES:
                    caption = f"{fig_key} {caption_tail}" if caption_tail else fig_key
                    insert_figure(doc, fig_key, caption)
                    i += 1
                    continue
            # 普通图注居中
            para = doc.add_paragraph()
            set_para_format(para, WD_ALIGN_PARAGRAPH.CENTER)
            process_inline_text(para, text, 10.5, bold=False, superscript_citations=not in_references)
            i += 1
            continue

        # 列表项 (- text) — 首行缩进0.74cm
        if line.strip().startswith("- "):
            text = line.strip()[2:]
            para = doc.add_paragraph()
            set_para_format(para, WD_ALIGN_PARAGRAPH.JUSTIFY,
                          first_line_indent_twips=TPL_LIST_INDENT)
            run = para.add_run("• ")
            set_run_font(run, 12, bold=False)
            process_inline_text(para, text, 12, bold=False, superscript_citations=not in_references)
            i += 1
            continue

        # 有序列表 (1. text) — 首行缩进0.74cm
        if re.match(r'^\d+\.\s', line.strip()):
            match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
            if match:
                para = doc.add_paragraph()
                set_para_format(para, WD_ALIGN_PARAGRAPH.JUSTIFY,
                              first_line_indent_twips=TPL_LIST_INDENT)
                run = para.add_run(f"{match.group(1)}. ")
                set_run_font(run, 12, bold=False)
                process_inline_text(para, match.group(2), 12, bold=False, superscript_citations=not in_references)
            i += 1
            continue

        # 普通段落 — 两端对齐, 无首行缩进
        # 跳过ASCII框图后残留的独立图注（已在代码块处理时插入）
        if re.match(r'^图\d+-\d+\s+.*$', line.strip()):
            i += 1
            continue
        para = doc.add_paragraph()
        set_para_format(para, WD_ALIGN_PARAGRAPH.LEFT if in_references else WD_ALIGN_PARAGRAPH.JUSTIFY)
        process_inline_text(para, line.strip(), 12, bold=False, superscript_citations=not in_references)
        i += 1

    doc.save(OUTPUT_PATH)
    print(f"DOCX已保存: {OUTPUT_PATH}")


if __name__ == "__main__":
    convert_md_to_docx()
